import streamlit as st
from docx import Document
import re
from io import BytesIO
import os

def clean_docx(file):
    doc = Document(file)

    # Patterns to remove: [H<number>], [hed], [dek], plus optional trailing space
    tag_patterns = [
        r'\[H\d+\]\s*',     # [H1], [H3], etc., plus space if any
        r'\[hed\]\s*',      # [hed] + space
        r'\[dek\]\s*'       # [dek] + space
    ]

    def clean_text(text):
        for pattern in tag_patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE)
        return text

    # Clean paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = clean_text(run.text)

    # Clean tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = clean_text(run.text)

    # Remove visible comment indicators
    doc_element = doc._element
    for comment in doc_element.xpath('//w:commentRangeStart | //w:commentRangeEnd | //w:commentReference'):
        comment.getparent().remove(comment)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

st.title("🧼 DOCX Cleaner")

uploaded_file = st.file_uploader("Upload a .docx file", type="docx")

if uploaded_file:
    cleaned_file = clean_docx(uploaded_file)
    
    base_filename = os.path.splitext(uploaded_file.name)[0]
    output_filename = f"{base_filename}_cleaned.docx"

    st.success("Cleaning complete!")
    st.download_button(
        label="Download cleaned .docx",
        data=cleaned_file,
        file_name=output_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
